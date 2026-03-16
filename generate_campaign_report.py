from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

import torch

from attenuate.model import architecture_summary


def _require_docx():
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:
        raise SystemExit("python-docx is required to generate the campaign report.") from exc
    return Document, Inches


def _read_json(path: str | Path) -> dict[str, Any]:
    return json.loads(Path(path).read_text(encoding="utf-8"))


def _format_metric(value: Any, digits: int = 4) -> str:
    if value is None:
        return "-"
    return f"{float(value):.{digits}f}"


def _resolve_report_model_cfg(summary: dict[str, Any]) -> dict[str, Any]:
    final_model = summary.get("final_model", {})
    if isinstance(final_model.get("model"), dict):
        return dict(final_model["model"])

    winner_name = summary.get("winner_direction")
    final_stage = final_model.get("stage")
    for direction in summary.get("directions", []):
        if direction.get("name") != winner_name:
            continue
        if final_stage == "stage2" and isinstance(direction.get("stage2"), dict):
            return dict(direction["stage2"].get("config", {}).get("model", {}))
        selected = direction.get("selected", {})
        if isinstance(selected.get("config"), dict):
            return dict(selected["config"].get("model", {}))
    return {"kind": "atennuate"}


def _architecture_paragraph(architecture: dict[str, Any]) -> str:
    if architecture["kind"] == "mp_senet_lite":
        return (
            "Arhitectura folosita este un denoiser time-frequency complex, phase-aware, "
            "de tip MP-SENet-lite, cu STFT interna, decodare paralela pentru magnitudine "
            "si faza, apoi reconstructie iSTFT."
        )
    return (
        "Arhitectura folosita este un denoiser aTENNuate-like pe waveform brut, mono 16 kHz, "
        "cu encoder-decoder pe blocuri SSM si inferenta offline."
    )


def _architecture_rows(architecture: dict[str, Any]) -> list[list[str]]:
    common = [
        ["Backend", architecture["kind"]],
        ["Sample rate", str(architecture["sample_rate"])],
        ["Parametri", f"{architecture['params']:,}"],
        ["Mod de lucru", architecture["mode"]],
    ]
    if architecture["kind"] == "mp_senet_lite":
        return common + [
            ["n_fft / hop / win", f"{architecture['n_fft']} / {architecture['hop_length']} / {architecture['win_length']}"],
            ["Compresie magnitudine", _format_metric(architecture["compress_factor"], 3)],
            ["Canale baza / bottleneck", f"{architecture['base_channels']} / {architecture['bottleneck_channels']}"],
            ["Blocuri TF / heads", f"{architecture['num_tf_blocks']} / {architecture['num_heads']}"],
            ["Dense growth / layers", f"{architecture['dense_growth']} / {architecture['dense_layers']}"],
            ["Padding multiplu", str(architecture["padding_multiple"])],
        ]
    return common + [
        ["Canale encoder", ", ".join(str(v) for v in architecture["channels"])],
        ["Factori resampling", ", ".join(str(v) for v in architecture["resample_factors"])],
        ["Num coeffs / repeat", f"{architecture['num_coeffs']} / {architecture['repeat']}"],
        ["Padding multiplu", str(architecture["padding_multiple"])],
        ["Pre-conv", str(architecture["pre_conv"])],
    ]


def _add_table(document: Any, title: str, headers: list[str], rows: list[list[str]]) -> None:
    document.add_heading(title, level=2)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def generate_campaign_report(summary_path: str | Path, out_path: str | Path) -> None:
    Document, Inches = _require_docx()
    summary = _read_json(summary_path)
    architecture = architecture_summary(_resolve_report_model_cfg(summary))
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.add_heading("VoiceBank PESQ-First Campaign Report", level=0)
    document.add_paragraph(
        "Raport generat automat pentru campania end-to-end pe VoiceBank-DEMAND. "
        "Metrica primara de selectie este PESQ, cu guardrails pe STOI si SI-SDR."
    )

    dataset = summary.get("dataset", {})
    dataset_validation = dataset.get("validation", {})
    final_model = summary.get("final_model", {})
    exports = summary.get("exports", {})

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(
        f"Modelul final selectat: {final_model.get('name', '-')}, stage={final_model.get('stage', '-')}, "
        f"val PESQ={_format_metric(final_model.get('val', {}).get('PESQ'))}, "
        f"test PESQ={_format_metric(final_model.get('test', {}).get('PESQ'))}."
    )
    document.add_paragraph(
        f"Dataset validation: train={dataset_validation.get('counts', {}).get('train', '-')}, "
        f"val={dataset_validation.get('counts', {}).get('val', '-')}, "
        f"val_quick={dataset_validation.get('counts', {}).get('val_quick', '-')}, "
        f"test={dataset_validation.get('counts', {}).get('test', '-')}, "
        f"speaker_disjoint={dataset_validation.get('speaker_disjoint', False)}."
    )
    document.add_paragraph(
        f"Hardware/stack: device={summary.get('device', '-')}, VRAM={_format_metric(summary.get('gpu_memory_gb'), 2)} GB, "
        f"torch={torch.__version__}."
    )

    document.add_heading("Architecture", level=1)
    document.add_paragraph(_architecture_paragraph(architecture))
    _add_table(document, "Architecture Details", ["Field", "Value"], _architecture_rows(architecture))

    document.add_heading("Campaign Results", level=1)
    direction_rows: list[list[str]] = []
    for direction in summary.get("directions", []):
        base = direction.get("base", {})
        refined = direction.get("refined")
        selected = direction.get("selected", {})
        selected_test = direction.get("selected_test", {})
        analysis = direction.get("analysis", {})
        direction_rows.append(
            [
                direction.get("name", "-"),
                _format_metric(base.get("full_val", {}).get("PESQ")),
                _format_metric(refined.get("full_val", {}).get("PESQ")) if refined else "-",
                _format_metric(selected.get("full_val", {}).get("PESQ")),
                _format_metric(selected_test.get("PESQ")),
                analysis.get("rule", "-"),
                selected.get("stage", "-"),
            ]
        )
    _add_table(
        document,
        "Per-Direction Leaderboard",
        ["Direction", "Base Val PESQ", "Refined Val PESQ", "Best Val PESQ", "Best Test PESQ", "Rule", "Selected Stage"],
        direction_rows,
    )

    for direction in summary.get("directions", []):
        document.add_heading(direction.get("name", "-"), level=2)
        analysis = direction.get("analysis", {})
        base = direction.get("base", {})
        refined = direction.get("refined")
        selected = direction.get("selected", {})
        selected_test = direction.get("selected_test", {})
        issues = ", ".join(analysis.get("issues", [])) or "niciun issue critic"
        document.add_paragraph(
            f"Ipoteza/refinement: rule={analysis.get('rule', '-')}; rationale={analysis.get('rationale', '-')}; "
            f"issues={issues}."
        )
        document.add_paragraph(
            f"Base val: PESQ={_format_metric(base.get('full_val', {}).get('PESQ'))}, "
            f"STOI={_format_metric(base.get('full_val', {}).get('STOI'))}, "
            f"SI-SDR={_format_metric(base.get('full_val', {}).get('SI_SDR'))}."
        )
        if refined:
            document.add_paragraph(
                f"Refined val: PESQ={_format_metric(refined.get('full_val', {}).get('PESQ'))}, "
                f"STOI={_format_metric(refined.get('full_val', {}).get('STOI'))}, "
                f"SI-SDR={_format_metric(refined.get('full_val', {}).get('SI_SDR'))}."
            )
        document.add_paragraph(
            f"Best selected ({selected.get('stage', '-')}): val PESQ={_format_metric(selected.get('full_val', {}).get('PESQ'))}, "
            f"test PESQ={_format_metric(selected_test.get('PESQ'))}. "
            f"Checkpoint={selected.get('summary', {}).get('best_checkpoint', '-')}."
        )
        plot_paths = selected.get("summary", {}).get("plot_paths", {})
        metric_plot = plot_paths.get("metric_curves")
        if metric_plot and Path(metric_plot).exists():
            document.add_picture(metric_plot, width=Inches(6.2))

    document.add_heading("Final Model", level=1)
    document.add_paragraph(
        f"Winner direction: {summary.get('winner_direction', '-')}. "
        f"Final checkpoint: {final_model.get('checkpoint', '-')}. "
        f"Stage2 applied={summary.get('notes', {}).get('stage2_applied', False)}; "
        f"stable_direction={summary.get('notes', {}).get('stable_direction', False)}."
    )
    document.add_paragraph(
        f"Final validation metrics: PESQ={_format_metric(final_model.get('val', {}).get('PESQ'))}, "
        f"STOI={_format_metric(final_model.get('val', {}).get('STOI'))}, "
        f"SI-SDR={_format_metric(final_model.get('val', {}).get('SI_SDR'))}."
    )
    document.add_paragraph(
        f"Final test metrics: PESQ={_format_metric(final_model.get('test', {}).get('PESQ'))}, "
        f"STOI={_format_metric(final_model.get('test', {}).get('STOI'))}, "
        f"SI-SDR={_format_metric(final_model.get('test', {}).get('SI_SDR'))}, "
        f"Delta-SNR={_format_metric(final_model.get('test', {}).get('DELTA_SNR'))}."
    )

    document.add_heading("Exports", level=1)
    export_rows = [
        [
            "TorchScript",
            str(exports.get("torchscript", {}).get("success", False)),
            exports.get("torchscript", {}).get("artifact", "-"),
            exports.get("torchscript", {}).get("stderr", "")[-200:],
        ],
        [
            "ONNX",
            str(exports.get("onnx", {}).get("success", False)),
            exports.get("onnx", {}).get("artifact", "-"),
            exports.get("onnx", {}).get("stderr", "")[-200:] or exports.get("onnx", {}).get("stdout", "")[-200:],
        ],
    ]
    _add_table(document, "Export Status", ["Format", "Success", "Artifact", "Notes"], export_rows)

    document.add_heading("Reproducibility", level=1)
    document.add_paragraph(
        f"Config source: {summary.get('config_path', '-')}. "
        "Outputs are aggregated in campaign_summary.json and leaderboard.csv/json under the run directory."
    )
    document.add_paragraph(
        "Seed principal: 1337. Seed confirmare: 2337. "
        "Fiecare directie a fost antrenata end-to-end; smoke testele nu intra in raportul de performanta."
    )

    document.save(out_path.as_posix())


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate DOCX report from campaign_summary.json.")
    parser.add_argument("--summary", required=True, help="Path to campaign_summary.json")
    parser.add_argument("--out", required=True, help="Path to output .docx file")
    args = parser.parse_args()
    generate_campaign_report(args.summary, args.out)


if __name__ == "__main__":
    main()
